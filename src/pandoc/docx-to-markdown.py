import re
import docx
import pypandoc


def convert_word_to_markdown(docx_file_path, output_file_path):
    # 读取Word文档
    doc = docx.Document(docx_file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    # 处理表格
    for table in doc.tables:
        if len(table.columns) == 1 and len(table.rows) == 1:
            cell_text = table.cell(0, 0).text.strip()
            # 将表格内容转换为Markdown格式
            markdown_text = f"`{cell_text}`"
            # 替换表格内容
            full_text = [re.sub(re.escape(cell_text), markdown_text, para) for para in full_text]
    # 将文本转换为Markdown格式
    markdown_text = pypandoc.convert_text("\n".join(full_text), "md", format="docx")
    # 将Markdown格式保存到文件中
    with open(output_file_path, "w") as f:
        f.write(markdown_text)


docx_file_path = r"F:\document\dev-docs\01_cpp\01_C\C01_搭建C语言开发环境.docx"
output_file_path = r"F:\\document\\dev-docs\\01_cpp\\01_C\code.md"
convert_word_to_markdown(docx_file_path, output_file_path)
